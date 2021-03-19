# This is a sample Python script.

# Press Shift+F10 to execute it or replace it with your code.
# Press Double Shift to search everywhere for classes, files, tool windows, actions, and settings.


# Press the green button in the gutter to run the script.
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docxtpl import DocxTemplate, InlineImage
from docxcompose.composer import Composer


def dowork(filepath):
    allfiles = os.listdir(filepath)
    # 遍历该文件夹下的所有文件夹
    for file in allfiles:
        tem_path = os.path.join(filepath, file)
        # print(file == '三改一拆')
        # 判断该文件夹类型
        if file == '三改一拆':
            print(file)
            # 遍历得到所有年份目录
            for year_path in getdir(tem_path):
                print('获取年份目录 ：' + year_path)
                docname = year_path + '年' + file + '台账' + '.docx'
                document = Document()
                target_composer = Composer(document)
                # 遍历年份目录得到所有月份目录
                for month_path in getdir(os.path.join(filepath, file, year_path)):
                    print('获取月份目录 ：' + month_path)
                    # 遍历所有月份目录，得到每月台账资料
                    for mc_path in getdir(os.path.join(filepath, file, year_path, month_path)):
                        print('获取每月台账资料 ：' + mc_path)
                        piclist = []
                        txtlist = []
                        for mc_file in os.listdir(os.path.join(filepath, file, year_path, month_path, mc_path)):
                            print('获取每项台账资料文件 ：' + mc_file)
                            if os.path.splitext(mc_file)[1] == '.txt':
                                txtlist.append(mc_file)
                            else:
                                piclist.append(os.path.join(filepath, file, year_path,month_path, mc_path, mc_file))
                        if txtlist.__len__() > 1:
                            print('there are more than one txt file in ' + os.path.join(filepath, file, year_path,
                                                                                        month_path, mc_path))
                        else:
                            tem_docx = DocxTemplate("template1.docx")
                            result_content_dict = {}
                            count = 0
                            while count < piclist.__len__():
                                image = InlineImage(tem_docx, piclist[count], width=Cm(7), height=Cm(5))
                                result_content_dict["pic" + str(count)] = image
                                count += 1
                            for line in open(os.path.join(filepath, file, year_path, month_path, mc_path, txtlist[0]),
                                             "r", encoding="utf-8"):  # 设置文件对象并读取每一行文件
                                line = line.strip("\n")
                                data_list = line.split("#", 1)
                                if len(data_list):
                                    result_content_dict[data_list[0]] = data_list[1]
                            print(result_content_dict)
                            tem_docx.render(result_content_dict)
                            tem_docx.add_page_break()
                            target_composer.append(tem_docx)
                            # # 添加标题
                            # title = document.add_paragraph()
                            # run = title.add_run("三改一拆工作")
                            # run.font.bold = True  # 加粗
                            # # title.font.italic = True  # 斜体
                            # # title.font.underline = True  # 下划线
                            # # title.font.strike = True  # 删除线
                            # # title.font.shadow = True  # 阴影
                            # run.font.size = Pt(16)  # 字体大小
                            # run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色
                            # run.font.name = "仿宋GB2312"  # 字体
                            # run.element.rPr.rFonts.set(qn('w:eastAsia'), "仿宋GB2312")  # 中文字体
                            # title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            # # 添加标题
                            #
                            # # 添加表格
                            # table = document.add_table(rows=5, cols=3, style="Normal Table")
                            # table.alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            # table.cell(1, 1).width = Cm(3)
                            # table.cell(1, 1).height = Cm(0.6)
                            #
                            # document.add_paragraph("日期" + concentdict["日期"])
                target_composer.save(os.path.join(filepath, file, year_path) + "\\" + docname)
        elif file == '安全生产':
            print(file)
        else:
            print("无此类型台账模板：" + file)
        # if os.path.isdir(tem_path):
        #     getDir(tem_path)
        # 如果不是文件夹，保存文件路径及文件名
        # elif os.path.isfile(tem_path):
        #     allpath.append(tem_path)
        #     allname.append(file)


def getdir(filepath):
    allpath = []
    allfiles = os.listdir(filepath)
    # 遍历该文件夹下的所有文件夹
    for file in allfiles:
        tem_path = os.path.join(filepath, file)
        if os.path.isdir(tem_path):
            allpath.append(file)
    return allpath


def getcontent(txtfile):
    dict = {}
    for line in open(txtfile, "r", encoding="utf-8"):  # 设置文件对象并读取每一行文件
        line = line.strip("\n")
        datalist = line.split("#", 1)
        if len(datalist):
            dict[datalist[0]] = datalist[1]
    return dict


if __name__ == "__main__":
    path = "D:\\testpackage"
    folder = os.path.exists(path)
    if not folder:
        print("路径不存在")
    else:
        dowork(path)
    # for file in allpath:
    #     print(file)
    # print("-------------------------")
    # for name in allname:
    #     print(name)

# def traverse(f):
#     fs = os.listdir(f)
#     for f1 in fs:
#         tmp_path = os.path.join(f, f1)
#         if not os.path.isdir(tmp_path):
#             print('文件: %s' % tmp_path)
#         else:
#             print('文件夹：%s' % tmp_path)
#             traverse(tmp_path)
#
#
# path = 'E:/3、台账专用包/原桌面文件夹/执法业务台账'
# traverse(path)


# See PyCharm help at https://www.jetbrains.com/help/pycharm/
